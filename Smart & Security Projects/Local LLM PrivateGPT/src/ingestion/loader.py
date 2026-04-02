"""Multi-format document loader — PDF, DOCX, TXT, CSV, Markdown."""

import os
from dataclasses import dataclass, field
from typing import List

import fitz  # PyMuPDF
from docx import Document as DocxDocument
import pandas as pd


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------
@dataclass
class DocumentPage:
    """A single page / section extracted from a source file."""
    text: str
    page_number: int = 0
    metadata: dict = field(default_factory=dict)


@dataclass
class LoadedDocument:
    """Full document with list of pages and file-level metadata."""
    pages: List[DocumentPage]
    filename: str = ""
    file_type: str = ""
    file_size: int = 0
    page_count: int = 0


# ---------------------------------------------------------------------------
# Magic-byte validation helpers
# ---------------------------------------------------------------------------
_MAGIC_BYTES = {
    "pdf": b"%PDF",
    "docx": b"PK",       # ZIP archive (OOXML)
    "csv": None,          # no reliable magic bytes
    "txt": None,
    "md": None,
}

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "md"}


def _validate_magic(file_path: str, ext: str) -> bool:
    """Return True if the file's leading bytes match the expected signature."""
    expected = _MAGIC_BYTES.get(ext)
    if expected is None:
        return True  # no signature to check
    with open(file_path, "rb") as f:
        header = f.read(len(expected))
    return header == expected


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------
def _load_pdf(file_path: str) -> List[DocumentPage]:
    pages: List[DocumentPage] = []
    doc = fitz.open(file_path)
    for i, page in enumerate(doc):
        text = page.get_text("text")
        if text.strip():
            pages.append(DocumentPage(text=text, page_number=i + 1))
    doc.close()
    return pages


def _load_docx(file_path: str) -> List[DocumentPage]:
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Treat tables as additional text blocks
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)
    full_text = "\n".join(paragraphs)
    if full_text.strip():
        return [DocumentPage(text=full_text, page_number=1)]
    return []


def _load_txt(file_path: str) -> List[DocumentPage]:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    if text.strip():
        return [DocumentPage(text=text, page_number=1)]
    return []


def _load_markdown(file_path: str) -> List[DocumentPage]:
    """Load Markdown, splitting on top-level headers as separate pages."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    sections: List[DocumentPage] = []
    current_section = ""
    section_idx = 0
    for line in text.split("\n"):
        if line.startswith("# "):
            if current_section.strip():
                section_idx += 1
                sections.append(DocumentPage(text=current_section.strip(), page_number=section_idx))
            current_section = line + "\n"
        else:
            current_section += line + "\n"
    if current_section.strip():
        section_idx += 1
        sections.append(DocumentPage(text=current_section.strip(), page_number=section_idx))
    return sections


def _load_csv(file_path: str) -> List[DocumentPage]:
    df = pd.read_csv(file_path)
    rows: List[str] = []
    columns = list(df.columns)
    for _, row in df.iterrows():
        parts = [f"{col}: {row[col]}" for col in columns]
        rows.append(" | ".join(parts))
    text = "\n".join(rows)
    if text.strip():
        return [DocumentPage(text=text, page_number=1)]
    return []


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
_LOADERS = {
    "pdf": _load_pdf,
    "docx": _load_docx,
    "txt": _load_txt,
    "md": _load_markdown,
    "csv": _load_csv,
}


class DocumentLoader:
    """Load a supported document file and return structured text with metadata."""

    @staticmethod
    def load(file_path: str) -> LoadedDocument:
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lstrip(".").lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: .{ext}")

        if not _validate_magic(file_path, ext):
            raise ValueError(f"File magic bytes do not match expected type: .{ext}")

        loader_fn = _LOADERS[ext]
        pages = loader_fn(file_path)

        return LoadedDocument(
            pages=pages,
            filename=os.path.basename(file_path),
            file_type=ext,
            file_size=os.path.getsize(file_path),
            page_count=len(pages),
        )
